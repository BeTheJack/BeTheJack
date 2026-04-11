import streamlit as st
from groq import Groq
from fpdf import FPDF
import os
import re
import json
import io
from PIL import Image, ImageOps, ImageDraw

# ==============================================================================
# 1. SETUP & CONFIGURATION
# ==============================================================================
st.set_page_config(page_title="BeTheJack", page_icon="🃏", layout="wide")

# Custom CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Space Grotesk', sans-serif;
    }

    .main-header {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e);
        padding: 2rem 2.5rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        text-align: center;
    }
    .main-header h1 {
        color: #fff;
        font-size: 3rem;
        font-weight: 700;
        margin: 0;
        letter-spacing: -1px;
    }
    .main-header p {
        color: rgba(255,255,255,0.6);
        font-size: 1.1rem;
        margin: 0.3rem 0 0 0;
    }

    .step-badge {
        display: inline-block;
        background: linear-gradient(135deg, #667eea, #764ba2);
        color: white;
        border-radius: 50%;
        width: 28px;
        height: 28px;
        line-height: 28px;
        text-align: center;
        font-weight: 700;
        font-size: 0.85rem;
        margin-right: 8px;
    }

    .info-box {
        background: #f0f4ff;
        border-left: 4px solid #667eea;
        padding: 0.8rem 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
        font-size: 0.9rem;
        color: #334;
    }

    .success-box {
        background: #f0fff4;
        border-left: 4px solid #38a169;
        padding: 0.8rem 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
        font-size: 0.9rem;
        color: #276749;
    }

    .warning-box {
        background: #fffbeb;
        border-left: 4px solid #f59e0b;
        padding: 0.8rem 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
        font-size: 0.9rem;
        color: #78350f;
    }

    .stButton>button {
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.2s;
    }

    .diff-added { color: #16a34a; background: #f0fdf4; padding: 2px 4px; border-radius: 3px; }
    .diff-changed { color: #d97706; background: #fffbeb; padding: 2px 4px; border-radius: 3px; }
</style>
""", unsafe_allow_html=True)


def init_ai():
    import os
    try:
        api_key = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        api_key = ""
    if not api_key:
        api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        st.error("🚨 API Key Missing. Set GROQ_API_KEY in environment variables or Streamlit Secrets.")
        return False
    try:
        client = Groq(api_key=api_key)
        st.session_state["_groq_client"] = client
        return True
    except Exception as e:
        st.error(f"🚨 AI init failed: {e}")
        return False


# ==============================================================================
# 2. CV PARSING — Extract real data from uploaded CV
# ==============================================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (best for layout) with pypdf fallback."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception:
        pass

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    except Exception as e:
        st.error(f"Could not parse PDF: {e}")
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        st.error(f"Could not parse DOCX: {e}")
        return ""


def parse_cv_file(uploaded_file) -> str:
    """Route to the correct parser based on file type."""
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        st.warning("Unsupported file type. Please upload PDF, DOCX, or TXT.")
        return ""


# ==============================================================================
# 3. AI — Tailor CV content to JD
# ==============================================================================

def get_groq_client():
    """Retrieve the Groq client — falls back to creating a new one if session state missing."""
    import os
    client = st.session_state.get("_groq_client")
    if client:
        return client
    key = ""
    try:
        key = st.secrets.get("GROQ_API_KEY", "")
    except Exception:
        pass
    if not key:
        key = os.environ.get("GROQ_API_KEY", "")
    return Groq(api_key=key) if key else None


def get_best_model() -> str:
    """
    Returns the best available Groq model for CV generation.
    Priority: llama-3.3-70b (best quality) → llama-3.1-70b → llama3-70b → 8b fallback.
    """
    client = get_groq_client()
    if not client:
        return "llama-3.3-70b-versatile"
    try:
        available = [m.id for m in client.models.list().data]
        priorities = [
            "llama-3.3-70b-versatile",
            "llama-3.1-70b-versatile",
            "llama3-70b-8192",
            "llama-3.1-8b-instant",
            "llama3-8b-8192",
        ]
        for p in priorities:
            if p in available:
                return p
        for m in available:
            if "llama" in m.lower():
                return m
    except Exception:
        pass
    return "llama-3.3-70b-versatile"


def enforce_bullet_limit(text: str, max_bullets: int = 3) -> str:
    """
    Enforces bullet caps per role block.
    - Sub-role 2-part pipe lines (under ##COMPANY##): cap 3.
    - Normal 3-part pipe flat roles: cap 3.
    - 2-part pipe project lines: cap 1.
    - Blank lines reset the counter between roles.
    """
    lines = text.split('\n')
    result = []
    bullet_count = 0
    current_cap  = max_bullets

    for line in lines:
        stripped = line.strip()

        if not stripped:
            bullet_count = 0
            result.append(line)
            continue

        if stripped.startswith('##COMPANY##'):
            bullet_count = 0
            current_cap  = max_bullets
            result.append(line)
            continue

        if '|' in stripped and not stripped.startswith('-'):
            parts = [p.strip() for p in stripped.split('|')]
            bullet_count = 0
            if len(parts) >= 3:
                current_cap = max_bullets
            elif len(parts) == 2:
                import re as _re
                is_date = bool(_re.search(
                    r'\d{4}|Present|present|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', parts[1]))
                current_cap = max_bullets if is_date else 1
            result.append(line)
            continue

        if (stripped.isupper() and len(stripped) > 3
                and '|' not in stripped
                and not any(c.isdigit() for c in stripped)):
            bullet_count = 0
            current_cap  = max_bullets
            result.append(line)
            continue

        if stripped.startswith('-') and len(stripped) > 1:
            bullet_count += 1
            if bullet_count > current_cap:
                continue
            result.append(line)
            continue

        result.append(line)

    return '\n'.join(result)


def fix_company_markers(text: str) -> str:
    """
    Converts any COBLOCK/COMPANY line to ##COMPANY## Name.
    Handles all variants Llama produces.
    """
    pattern = re.compile(
        r'^[#*\s]*(?:COBLOCK|COMPANY)[#*:\-\s]*(.+)$',
        re.IGNORECASE
    )
    SECTION_KEYWORDS = {
        'PROFESSIONAL EXPERIENCE', 'TECHNICAL SKILLS', 'SKILLS',
        'PROJECTS', 'EDUCATION', 'CERTIFICATIONS', 'CONTACT',
        'INTRODUCTION', 'SUMMARY', 'NAME'
    }
    result = []
    for line in text.split('\n'):
        s = line.strip()
        s_clean = re.sub(r'\*+', '', s).strip()
        m = pattern.match(s_clean)
        if m:
            name = m.group(1).strip().lstrip('#*:- ')
            if name and name.upper() not in SECTION_KEYWORDS:
                result.append(f'##COMPANY## {name}')
                continue
        if s.startswith('##COMPANY##'):
            name = s[len('##COMPANY##'):].strip()
            result.append(f'##COMPANY## {name}' if name else '')
            continue
        result.append(line)
    return '\n'.join(result)


def fix_bullet_markers(text: str) -> str:
    """
    Converts > bullet prefix from prompt to - prefix that the PDF builder expects.
    Also strips any leftover markdown bold/italic from bullet text.
    """
    result = []
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('> ') or stripped == '>':
            content = stripped[1:].lstrip()
            content = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', content)
            result.append('- ' + content)
        else:
            cleaned = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', line)
            result.append(cleaned)
    return '\n'.join(result)


def enforce_certifications(tailored: str, original_cv: str) -> str:
    """
    Strips any certification bullet from the tailored CV that doesn't appear
    (loosely) in the original CV text. Prevents hallucinated certs.
    If no certs exist in original, replaces section content with 'None listed.'
    """
    lines = tailored.split('\n')
    result = []
    in_certs = False
    original_lower = original_cv.lower()

    # Check if original CV has ANY certification content at all
    has_any_certs = bool(re.search(
        r'certif|cisco|comptia|aws|azure|google cloud|pmp|itil|scrum|prince2|ccna|ccnp|mcsa|mcse|rhce|ceh|oscp',
        original_lower
    ))

    cert_bullets_written = 0

    for line in lines:
        stripped = line.strip()

        if stripped.upper() == 'CERTIFICATIONS':
            in_certs = True
            cert_bullets_written = 0
            result.append(line)
            continue

        if in_certs:
            # Exit cert section on new ALL-CAPS section header
            if (stripped.isupper() and len(stripped) > 3
                    and '|' not in stripped
                    and not stripped.startswith('-')
                    and stripped.upper() != 'CERTIFICATIONS'):
                # If no real certs were written, add a placeholder
                if cert_bullets_written == 0:
                    result.append('- None listed')
                in_certs = False
                result.append(line)
                continue

            if stripped.startswith('-') and len(stripped) > 1:
                cert_text = stripped[1:].strip().lower()

                if not has_any_certs:
                    # Original CV has zero certs — drop everything
                    continue

                # Check if meaningful words from this cert appear in original CV
                words = [w for w in re.split(r'\W+', cert_text) if len(w) > 3]
                matched = sum(1 for w in words if w in original_lower)
                threshold = min(2, len(words))

                if len(words) == 0 or matched < threshold:
                    continue  # invented cert — drop it

                result.append(line)
                cert_bullets_written += 1
                continue

            # Non-bullet line inside certs (e.g. "None listed" plain text) — keep
            result.append(line)
            continue

        result.append(line)

    # Handle case where certs section was last and had no valid bullets
    if in_certs and cert_bullets_written == 0:
        result.append('- None listed')

    return '\n'.join(result)


def enforce_jake_format(text: str) -> str:
    """
    Cleans up common LLM format deviations for the India/Jake single-column layout.
    - Removes stray markdown (**, ##, ---)
    - Normalises section headers to ALL CAPS
    - Ensures skill lines use 'Category: items' not '- Category: items'
    - Removes [SIDEBAR_START] / [MAIN_START] markers that bleed into India output
    - Strips any stray pipe-less date lines that look like misformatted role headers
    """
    known_sections = {
        'professional experience', 'technical skills', 'skills',
        'projects', 'education', 'certifications', 'contact',
        'introduction', 'summary', 'name'
    }
    lines = text.split('\n')
    result = []

    for line in lines:
        # Strip sidebar markers
        line = line.replace('[SIDEBAR_START]', '').replace('[MAIN_START]', '')

        # Strip markdown bold/italic/headers
        line = re.sub(r'\*{1,3}([^*]*)\*{1,3}', r'\1', line)
        line = re.sub(r'^#{1,6}\s*', '', line)
        line = re.sub(r'^---+$', '', line)

        stripped = line.strip()
        if not stripped:
            result.append('')
            continue

        # Normalise near-match section headers to ALL CAPS
        lower = stripped.lower().rstrip(':').strip()
        if lower in known_sections:
            result.append(stripped.upper().rstrip(':'))
            continue

        # Fix skill lines that start with a dash but contain a colon
        # "- Category: items" → "Category: items"
        # Only unwrap if it looks like a skill (short category before colon)
        if stripped.startswith('-') and ':' in stripped:
            potential = stripped[1:].lstrip()
            colon_idx = potential.index(':')
            cat_candidate = potential[:colon_idx].strip()
            if len(cat_candidate.split()) <= 4 and len(cat_candidate) > 1:
                result.append(potential)
                continue

        result.append(line)

    return '\n'.join(result)


def tailor_cv(raw_cv_text: str, job_description: str, style: str = "Global") -> str:
    """
    Takes real extracted CV text and the target JD.
    Returns a tailored, enhanced version with up to 30% JD-driven augmentation.
    """
    client     = get_groq_client()
    model_name = get_best_model()

    visa_note = (
        "Extract Visa Status and Nationality from the CV if present; include in CONTACT section."
        if style == "Global"
        else "Do NOT include Visa Status, Nationality, or Photo."
    )

    layout_note = (
        "Use a two-column layout marker: write [SIDEBAR_START] before the sidebar section "
        "(Contact, Introduction, Skills, Certifications, Education) and [MAIN_START] before "
        "the main section (Experience, Projects)."
        if style == "Global"
        else "Single column, clean ATS-friendly layout. Do NOT include [SIDEBAR_START] or [MAIN_START] markers."
    )

    prompt = f"""
You are an experienced CV writer who specialises in taking thin, underdeveloped CVs and making them interview-worthy. Your job is to rewrite and expand the candidate's CV to match the target job — written like a real person, not an AI tool.

DETECTING A SHORT CV:
If the original CV has fewer than 3 bullets per role, vague one-liners, or roles that are under-described, you MUST expand them significantly. Do not just repeat what's there — infer the realistic day-to-day work someone in that role would do, and write it out properly. A Service Desk Engineer handles tickets, troubleshoots hardware/software, manages users, liaises with vendors, escalates issues, maintains documentation — write all of that out.

EXPANSION RULES:
1. Every role must have 3 bullets minimum, ideally 3 solid ones.
2. If a bullet is vague ("Resolved customer queries", "Troubleshooted hardware problems") — expand it into a proper sentence with real context: what kind of queries, on which systems, for how many users, what the process was.
3. If a role mentions a number (500 users, 1000 clients) — use that number but add the surrounding context that's missing.
4. For projects — if described in one vague line, expand to show what was actually built or done.
5. You can add 1-2 plausible bullets per role that a person in that position would genuinely do, even if not explicitly mentioned. Keep them believable and grounded.

BULLET WRITING STYLE:
Bullets should read like something you would say in an interview — specific, natural, not corporate.
Each bullet: one clear action + what/how + optional outcome. Aim for 15-22 words.
Outcomes do NOT have to be a percentage. Use variety:
  - Counts: "for around 200 users across 3 floors"
  - Timeframes: "resolved most tickets same-day"
  - Scope: "covering hardware, software and network issues"
  - Qualitative: "which meant the team stopped getting repeat calls on the same issue"
  - Only use % when the CV itself mentions it and it feels natural

MAX 1 percentage metric per role block. Do not put % in every bullet.

GOOD bullet examples:
> Handled first-line support for 200 users at HDFC Bank, covering Flexcube, Finone and Outlook 365 issues daily
> Set up and maintained antivirus across around 1000 endpoints at Wipro, running scheduled scans and dealing with any alerts that came through
> Configured network printers across multiple floors, fixed connectivity issues and kept a log of recurring problems to spot patterns
> Wrote a Python script to auto-categorise incoming ServiceNow tickets, which cut the manual triage time significantly
> Helped onboard new staff by setting up their laptops, configuring email, and walking them through the main systems they would be using

BAD bullet examples (avoid entirely):
> Resolved 500 customer queries, improved first-call resolution rate by 25%
> Troubleshooted 50 hardware problems, reduced resolution time by 30%
> Leveraged Python to streamline ticket automation processes

BANNED WORDS — never use:
spearheaded, leveraged, utilized, adept at, results-driven, dynamic, proven track record,
synergies, cutting-edge, innovative, streamlined, orchestrated, demonstrating, showcasing,
proficient, expertise in, robust, seamlessly, holistic, actionable, facilitated, ensured

INTRO STYLE:
Write it like the candidate typed it themselves. First sentence: where they've worked and for how long. Second: what they are specifically good at or enjoy.
Max 40 words. No buzzwords. No %.
GOOD: "I've spent 2 years in IT support at Wipro and HDFC Bank, mostly handling day-to-day desktop and application support. I'm good at getting to the root of software issues quickly and keeping users up and running."

COMPANY BLOCK FORMAT — use exactly this marker:
COBLOCK Reliance Industries - Mumbai

Sub-role under COBLOCK (title | dates, 2 parts):
Service Desk Engineer | 2024 - Present

Flat single role (3 parts):
Desktop Support Engineer | Wipro - Mumbai | 2022 - 11/2024

ALL bullets use > prefix. Never use - as bullet prefix.

RULES:
1. COBLOCK for companies with 2+ years or multiple roles. Flat 3-part for short single roles.
2. If only 1 title for 2+ years: invent 2-3 believable promotions, split dates proportionally.
3. Minimum 3 bullets per role. Max 3. Each 15-22 words. Expand vague lines.
4. MAX 1 % per role block. Use counts, scope and qualitative outcomes for the rest.
5. Projects: keep all real ones (expanded properly) + add exactly 2 new ones using JD tools.
6. Project bullets: 1 per project, 12-18 words, real specific description.
7. Skills: 4 categories, 4-5 items each. Add JD-adjacent tools plausibly.
8. Never invent companies, degrees, or certifications.
9. CERTIFICATIONS — CRITICAL RULE: Copy certifications EXACTLY and ONLY as they appear in the original CV. Do NOT invent, add, guess, or infer any certification not explicitly written in the original CV. If the original CV has zero certifications listed, write ONLY: CERTIFICATIONS followed by a new line with: - None listed. This rule is absolute — no exceptions.
10. {visa_note}
11. {layout_note}

SECTION ORDER (never change):
NAME > CONTACT > INTRODUCTION > TECHNICAL SKILLS > PROFESSIONAL EXPERIENCE > PROJECTS > EDUCATION > CERTIFICATIONS

FORMAT RULES:
- No ** bold anywhere
- No ### headers
- No --- dividers
- Section headers: ALL CAPS, no trailing colon
- Skill lines: Category: item1, item2, item3  (colon format, no > prefix, no leading dash)
- Project lines: Project Name | Tech1, Tech2  (2-part pipe, no dates)
- All bullets use > prefix

---
ORIGINAL CV:
{raw_cv_text}

---
TARGET JOB DESCRIPTION:
{job_description}

---
OUTPUT (write the actual content — no placeholders like [bullet] or [Project Name]):

NAME
[Full Name]

CONTACT
[Phone] | [Email] | [Location]

INTRODUCTION
[2 human sentences. Where worked + how long. What good at. Max 40 words. No %.]

TECHNICAL SKILLS
[Category]: [item1, item2, item3, item4]
[Category]: [item1, item2, item3, item4]
[Category]: [item1, item2, item3, item4]
[Category]: [item1, item2, item3, item4]

PROFESSIONAL EXPERIENCE

COBLOCK [Company - City]
[Most Recent Role] | [Start] - [End]
> [expanded bullet, 15-22 words]
> [expanded bullet, 15-22 words]
> [expanded bullet, 15-22 words]
[Previous Role] | [Start] - [End]
> [expanded bullet]
> [expanded bullet]
> [expanded bullet]

[Flat Title] | [Company - City] | [Start] - [End]
> [expanded bullet]
> [expanded bullet]
> [expanded bullet]

PROJECTS
[Real project name from CV] | [Tech Stack, enhanced with JD tools]
> [expanded specific description, 12-18 words]

[Invented project 1 — real name based on JD tools] | [JD tools]
> [what it does and outcome, 12-18 words]

[Invented project 2 — real name based on JD tools] | [JD tools]
> [what it does and outcome, 12-18 words]

EDUCATION
[Degree] | [University] | [Year]

CERTIFICATIONS
[Copy from original CV only. If none exist write: - None listed]
"""
    import time
    last_err = None
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an experienced CV writer. "
                            "Write naturally — like a real person, not an AI. "
                            "Never use corporate buzzwords. "
                            "Vary how you express impact: use counts, timeframes, team sizes and qualitative outcomes — "
                            "not a percentage on every bullet. Max 1 % per role block. "
                            "Never write placeholder text like [Project Name] — always write the actual content. "
                            "CERTIFICATIONS: copy only what is in the original CV. Never invent certifications. "
                            "If no certifications exist in the original CV, write: CERTIFICATIONS then on the next line: - None listed. "
                            "Output ONLY the resume. No commentary, no markdown, no explanations."
                        )
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.9,
                max_tokens=6000,
            )
            raw = response.choices[0].message.content
            raw = fix_company_markers(raw)
            raw = fix_bullet_markers(raw)
            if style == "India":
                raw = enforce_jake_format(raw)
            raw = enforce_certifications(raw, raw_cv_text)
            return enforce_bullet_limit(raw, max_bullets=3)
        except Exception as e:
            last_err = e
            err_str = str(e)
            if "429" in err_str or "rate" in err_str.lower() or "limit" in err_str.lower():
                wait = (attempt + 1) * 15
                st.warning(f"⏳ Rate limit hit — retrying in {wait}s... (attempt {attempt+1}/3)")
                time.sleep(wait)
            else:
                break
    return f"Error generating content: {last_err}"


# ==============================================================================
# 4. PDF BUILDER
# ==============================================================================


def sanitize(text: str) -> str:
    """
    Converts AI-generated Unicode text into FPDF-safe Latin-1 text.
    Maps every known problematic character explicitly before the final encode.
    Uses 'replace' as final fallback so it NEVER raises an exception.
    """
    replacements = {
        '\u2013': '-', '\u2014': '-', '\u2012': '-', '\u2015': '-', '\u2212': '-',
        '\u2018': "'", '\u2019': "'", '\u201a': "'", '\u201c': '"', '\u201d': '"',
        '\u201e': '"', '\u00ab': '"', '\u00bb': '"',
        '\u2022': '-', '\u2023': '-', '\u25cf': '-', '\u2219': '-', '\u00b7': '-',
        '\u2026': '...', '\u00a0': ' ', '\u202f': ' ', '\u2009': ' ',
        '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
        '\u2192': '->', '\u2190': '<-', '\u21d2': '=>',
        '\u00d7': 'x', '\u00f7': '/', '\u00b0': ' deg',
        '\u00ae': '(R)', '\u00a9': '(C)', '\u2122': '(TM)',
        '\u20ac': 'EUR', '\u00a3': 'GBP', '\u00a5': 'JPY',
        '\u00bd': '1/2', '\u00bc': '1/4', '\u00be': '3/4',
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    text = re.sub(r'\*{1,3}([^*]*)\*{1,3}', r'\1', text)
    text = re.sub(r'#{1,6}\s?', '', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    text = re.sub(r'`{1,3}[^`]*`{1,3}', '', text)

    text = text.encode('latin-1', errors='replace').decode('latin-1')
    return text


def crop_to_circle(image_path: str) -> str:
    try:
        img = Image.open(image_path).convert("RGB")
        size = min(img.size)
        img = ImageOps.fit(img, (size, size), centering=(0.5, 0.5))
        mask = Image.new('L', img.size, 0)
        ImageDraw.Draw(mask).ellipse((0, 0) + img.size, fill=255)
        img.putalpha(mask)
        out_path = "temp_circle_photo.png"
        img.save(out_path)
        return out_path
    except Exception:
        return image_path


class PDF(FPDF):
    pass


def build_docx(content: str) -> bytes:
    """
    Converts the tailored CV text into a clean, ATS-friendly DOCX file.
    Uses python-docx with professional formatting matching Jake's Resume style.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDocument()

    section = doc.sections[0]
    margin = Inches(0.75)
    section.top_margin    = margin
    section.bottom_margin = margin
    section.left_margin   = margin
    section.right_margin  = margin

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    pfmt = style.paragraph_format
    pfmt.space_after  = Pt(0)
    pfmt.space_before = Pt(0)

    def add_para(text='', bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=0, space_after=2, color=None, indent_left=0):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after  = Pt(space_after)
        if indent_left:
            pf.left_indent = Inches(indent_left)
        if text:
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.name = 'Calibri'
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_rule(color_hex='000000'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_two_col_run(p, left_text, left_bold, right_text, right_italic=True):
        p.clear()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)
        run_l = p.add_run(left_text)
        run_l.bold = left_bold
        run_l.font.size = Pt(10.5)
        run_l.font.name = 'Calibri'
        run_tab = p.add_run('\t')
        run_tab.font.name = 'Calibri'
        run_r = p.add_run(right_text)
        run_r.italic = right_italic
        run_r.font.size = Pt(9.5)
        run_r.font.name = 'Calibri'
        run_r.font.color.rgb = RGBColor(100, 100, 100)

    text = content.replace('[SIDEBAR_START]', '').replace('[MAIN_START]', '')
    lines = [l.rstrip() for l in text.split('\n')]

    SKIP_WORDS = {'NAME', 'CONTACT', 'SIDEBAR_START', 'MAIN_START'}
    NAVY = (18, 40, 76)

    name_line = next((l.strip() for l in lines
                      if l.strip() and l.strip() not in SKIP_WORDS
                      and l.strip() != 'INTRODUCTION'
                      and '|' not in l.strip()
                      and not l.strip().startswith('-')), '')
    contact_line = next((l.strip() for l in lines
                         if ('@' in l or l.strip().startswith('+') or
                             re.search(r'\d{7,}', l))
                         and l.strip() != name_line), '')

    name_done    = False
    contact_done = False
    intro_mode   = False
    in_experience = False

    for line in lines:
        line = line.strip()

        if line in SKIP_WORDS:
            continue

        if not line:
            continue

        if not name_done and line == name_line:
            p = add_para(line.title(), bold=True, size=20,
                         align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=0, space_after=2)
            name_done = True
            continue

        if name_done and not contact_done and line == contact_line:
            add_para(line, italic=False, size=9,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=(80, 80, 80), space_after=3)
            contact_done = True
            continue

        if line == 'INTRODUCTION':
            intro_mode = True
            continue

        if intro_mode:
            if (line.isupper() and len(line) > 3 and '|' not in line) or line.startswith('##'):
                intro_mode = False
            else:
                add_para(line, italic=True, size=9, color=(60, 60, 60), space_after=1)
                continue

        if line.startswith('##COMPANY##'):
            company_name = line.replace('##COMPANY##', '').strip()
            in_experience = True
            p = add_para(company_name.upper(), bold=True, size=11,
                         color=NAVY, space_before=8, space_after=0)
            add_rule('000000')
            continue

        if (line.isupper() and 3 < len(line) <= 35
                and '|' not in line and line not in SKIP_WORDS
                and not any(c.isdigit() for c in line)):
            in_experience = 'EXPERIENCE' in line
            p = add_para(line, bold=True, size=10.5,
                         color=NAVY, space_before=8, space_after=0)
            add_rule('000000')
            continue

        if '|' in line and not line.startswith('-'):
            parts = [p.strip() for p in line.split('|')]

            if len(parts) >= 3:
                title, company, dates = parts[0], parts[1], parts[2]
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(0)
                add_two_col_run(p, title, True, dates)
                add_para(company, italic=True, size=9.5, color=(60, 60, 60), space_after=1)

            elif len(parts) == 2:
                p0, p1 = parts[0], parts[1]
                is_date = bool(re.search(
                    r'\d{4}|Present|present|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', p1))
                if is_date:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(0)
                    p.paragraph_format.left_indent  = Inches(0.15)
                    add_two_col_run(p, p0, True, p1)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(0)
                    add_two_col_run(p, p0, True, p1)
            continue

        if ':' in line and not line.startswith('-'):
            colon_idx = line.index(':')
            cat = line[:colon_idx].strip()
            det = line[colon_idx + 1:].strip()
            if cat and det and len(cat.split()) <= 5:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(1)
                run_cat = p.add_run(cat + ': ')
                run_cat.bold = True
                run_cat.font.size = Pt(9.5)
                run_cat.font.name = 'Calibri'
                run_det = p.add_run(det)
                run_det.font.size = Pt(9.5)
                run_det.font.name = 'Calibri'
                continue

        if line.startswith('-'):
            content_text = line[1:].lstrip()
            indent = 0.25 if in_experience else 0.15
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(indent)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            run = p.add_run('\u2022  ' + content_text)
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            continue

        add_para(line, size=9.5, color=(50, 50, 50), space_after=1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf(content: str, style: str, photo_path: str = None, photo_size: int = 52) -> bytes:
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    content = sanitize(content)

    if style == "Global":
        _build_global_pdf(pdf, content, photo_path, photo_size=photo_size)
    else:
        _build_india_pdf(pdf, content)

    raw = pdf.output(dest='S')
    if isinstance(raw, (bytes, bytearray)):
        return bytes(raw)
    return raw.encode('latin-1', errors='replace')


def _build_global_pdf(pdf: FPDF, text: str, photo_path: str = None, photo_size: int = 52):
    """
    Premium two-column sidebar layout.
    Dark navy sidebar, gold accent line, circle photo, bold company names in main.
    """
    SIDEBAR_W    = 72
    SIDEBAR_X    = 6
    SIDEBAR_TW   = SIDEBAR_W - SIDEBAR_X - 4
    MAIN_X       = SIDEBAR_W + 6
    PAGE_W       = 210
    RIGHT_MARGIN = 9
    MAIN_W       = PAGE_W - MAIN_X - RIGHT_MARGIN
    PAGE_H       = 297

    NAVY         = (18,  40,  76)
    GOLD         = (180, 148,  80)
    SIDEBAR_TEXT = (220, 220, 220)
    SIDEBAR_DIM  = (160, 160, 160)
    WHITE        = (255, 255, 255)
    BLACK        = (10,  10,  10)
    DARK_GREY    = (50,  50,  50)
    MID_GREY     = (100, 100, 100)

    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, SIDEBAR_W, PAGE_H, 'F')

    pdf.set_fill_color(*GOLD)
    pdf.rect(0, 0, PAGE_W, 3, 'F')

    sidebar_text = ""
    main_text    = text
    if "[SIDEBAR_START]" in text and "[MAIN_START]" in text:
        parts        = text.split("[MAIN_START]", 1)
        sidebar_text = parts[0].replace("[SIDEBAR_START]", "").strip()
        main_text    = parts[1].strip()
    elif "PROFESSIONAL EXPERIENCE" in text:
        idx          = text.find("PROFESSIONAL EXPERIENCE")
        sidebar_text = text[:idx].strip()
        main_text    = text[idx:].strip()

    SKIP_WORDS = {"NAME", "CONTACT", "INTRODUCTION", "SIDEBAR_START", "MAIN_START"}

    cur_y = 8

    if photo_path and os.path.exists(photo_path):
        try:
            img  = Image.open(photo_path).convert("RGBA")
            size = min(img.size)
            img  = ImageOps.fit(img, (size, size), centering=(0.5, 0.25))
            mask = Image.new('L', (size, size), 0)
            ImageDraw.Draw(mask).ellipse((0, 0, size, size), fill=255)
            circle = Image.new('RGBA', (size, size), (0, 0, 0, 0))
            circle.paste(img, mask=mask)
            bg = Image.new('RGB', (size, size), (18, 40, 76))
            bg.paste(circle, mask=circle.split()[3])
            circ_path = "temp_circle_photo.png"
            bg.save(circ_path, "PNG")
            photo_x = (SIDEBAR_W - photo_size) / 2
            pdf.image(circ_path, x=photo_x, y=cur_y, w=photo_size)
            cur_y += photo_size + 4
            try: os.remove(circ_path)
            except: pass
        except Exception:
            cur_y = 8

    pdf.set_draw_color(*GOLD)
    pdf.set_line_width(0.6)
    pdf.line(SIDEBAR_X, cur_y, SIDEBAR_W - 4, cur_y)
    pdf.set_line_width(0.2)
    cur_y += 4

    pdf.set_xy(SIDEBAR_X, cur_y)

    name_line = next(
        (l.strip() for l in sidebar_text.split('\n')
         if l.strip() and l.strip() not in SKIP_WORDS and '|' not in l),
        ''
    )

    sidebar_intro_lines = 0

    for raw_line in sidebar_text.split('\n'):
        line = raw_line.strip()
        if pdf.get_y() > PAGE_H - 8:
            break
        if not line:
            pdf.set_xy(SIDEBAR_X, min(pdf.get_y() + 1.0, PAGE_H - 8))
            continue
        if line in SKIP_WORDS:
            continue

        if line == name_line:
            pdf.set_x(SIDEBAR_X)
            pdf.set_font("Arial", 'B', 13)
            pdf.set_text_color(*WHITE)
            pdf.multi_cell(SIDEBAR_TW, 6.5, line.upper(), align='C')
            uy = pdf.get_y() + 1
            pdf.set_draw_color(*GOLD)
            pdf.set_line_width(0.5)
            pdf.line(SIDEBAR_X + 4, uy, SIDEBAR_W - 8, uy)
            pdf.set_line_width(0.2)
            pdf.ln(3)
            continue

        if (line.isupper() and 3 < len(line) < 30
                and line not in SKIP_WORDS
                and not any(c.isdigit() for c in line)):
            sidebar_intro_lines = 0
            pdf.ln(3)
            pdf.set_x(SIDEBAR_X)
            pdf.set_font("Arial", 'B', 7.5)
            pdf.set_text_color(*GOLD)
            pdf.cell(SIDEBAR_TW, 4.5, line, ln=True)
            ry = pdf.get_y()
            pdf.set_draw_color(*GOLD)
            pdf.set_line_width(0.4)
            pdf.line(SIDEBAR_X, ry, SIDEBAR_W - 4, ry)
            pdf.set_line_width(0.2)
            pdf.ln(1.5)
            continue

        if ":" in line and not line.startswith("-"):
            colon_idx = line.index(":")
            cat = line[:colon_idx].strip()
            det = line[colon_idx + 1:].strip()
            if cat and det and len(cat.split()) <= 4:
                pdf.set_x(SIDEBAR_X)
                pdf.set_font("Arial", 'B', 7)
                pdf.set_text_color(*GOLD)
                lw = min(pdf.get_string_width(cat + ": ") + 1, SIDEBAR_TW - 4)
                pdf.cell(lw, 4, cat + ": ", ln=0)
                pdf.set_font("Arial", '', 7)
                pdf.set_text_color(*SIDEBAR_TEXT)
                pdf.multi_cell(SIDEBAR_TW - lw, 4, det, align='L')
                pdf.ln(0.2)
                continue

        if line.startswith("-") and ":" in line:
            cat_part, _, det_part = line.partition(":")
            cat = cat_part.replace("-", "").strip()
            det = det_part.strip()
            pdf.set_x(SIDEBAR_X)
            pdf.set_font("Arial", 'B', 7)
            pdf.set_text_color(*GOLD)
            label_w = min(pdf.get_string_width(cat + ": ") + 1, SIDEBAR_TW - 6)
            pdf.cell(label_w, 4, cat + ": ", ln=0)
            pdf.set_font("Arial", '', 7)
            pdf.set_text_color(*SIDEBAR_TEXT)
            pdf.multi_cell(SIDEBAR_TW - label_w, 4, det, align='L')
            pdf.ln(0.2)
            continue

        if line.startswith("-"):
            pdf.set_x(SIDEBAR_X + 2)
            pdf.set_font("Arial", '', 7.5)
            pdf.set_text_color(*SIDEBAR_TEXT)
            pdf.multi_cell(SIDEBAR_TW - 2, 4, "\x95 " + line[1:].lstrip(), align='L')
            continue

        if sidebar_intro_lines < 2:
            display = line[:120] + ('...' if len(line) > 120 else '')
            pdf.set_x(SIDEBAR_X)
            pdf.set_font("Arial", size=7.5)
            pdf.set_text_color(*SIDEBAR_DIM)
            pdf.multi_cell(SIDEBAR_TW, 4, display, align='L')
            sidebar_intro_lines += 1
        else:
            if len(line) < 50 or '|' in line or '@' in line or any(c.isdigit() for c in line[:8]):
                pdf.set_x(SIDEBAR_X)
                pdf.set_font("Arial", size=7.5)
                pdf.set_text_color(*SIDEBAR_DIM)
                pdf.multi_cell(SIDEBAR_TW, 4, line, align='L')

    # MAIN COLUMN
    pdf.set_xy(MAIN_X, 8)
    pdf.set_text_color(*BLACK)

    def _new_page():
        pdf.add_page()
        pdf.set_fill_color(*NAVY)
        pdf.rect(0, 0, SIDEBAR_W, PAGE_H, 'F')
        pdf.set_fill_color(*GOLD)
        pdf.rect(0, 0, PAGE_W, 3, 'F')
        pdf.set_xy(MAIN_X, 8)

    in_exp_global = False

    for raw_line in main_text.split('\n'):
        line = raw_line.strip()
        if not line:
            if pdf.get_y() < PAGE_H - 6:
                pdf.set_xy(MAIN_X, pdf.get_y() + 1.5)
            continue
        if pdf.get_y() > PAGE_H - 12:
            _new_page()

        if (line.isupper() and len(line) < 40
                and line not in SKIP_WORDS
                and not any(c.isdigit() for c in line)):
            in_exp_global = ("EXPERIENCE" in line)
            pdf.ln(5)
            pdf.set_x(MAIN_X)
            pdf.set_font("Arial", 'B', 11)
            pdf.set_text_color(*NAVY)
            pdf.cell(MAIN_W, 6, line, ln=True)
            ry = pdf.get_y()
            pdf.set_draw_color(*GOLD)
            pdf.set_line_width(0.5)
            pdf.line(MAIN_X, ry, MAIN_X + MAIN_W, ry)
            pdf.set_line_width(0.2)
            pdf.set_text_color(*BLACK)
            pdf.ln(3)
            continue

        if line.startswith("##COMPANY##"):
            company_name = line.replace("##COMPANY##", "").strip()
            if pdf.get_y() > PAGE_H - 18:
                _new_page()
            else:
                pdf.ln(4)
            pdf.set_x(MAIN_X)
            pdf.set_font("Arial", 'B', 11)
            pdf.set_text_color(*NAVY)
            pdf.cell(MAIN_W, 5.5, company_name.upper(), ln=True)
            ry = pdf.get_y()
            pdf.set_draw_color(*GOLD)
            pdf.set_line_width(0.4)
            pdf.line(MAIN_X, ry, MAIN_X + MAIN_W * 0.6, ry)
            pdf.set_line_width(0.2)
            pdf.set_text_color(*BLACK)
            pdf.ln(1.5)
            continue

        if "|" in line and not line.startswith("-"):
            parts = [p.strip() for p in line.split("|")]
            pdf.ln(3)
            pdf.set_x(MAIN_X)

            if len(parts) >= 3:
                title, company, dates = parts[0], parts[1], parts[2]
                pdf.set_font("Arial", 'B', 11)
                pdf.set_text_color(*NAVY)
                pdf.cell(MAIN_W, 5.5, company.upper(), ln=True)
                pdf.set_x(MAIN_X)
                pdf.set_font("Arial", 'B', 9.5)
                pdf.set_text_color(*DARK_GREY)
                title_w = MAIN_W * 0.65
                pdf.cell(title_w, 4.5, title, ln=0)
                pdf.set_font("Arial", 'I', 9)
                pdf.set_text_color(*MID_GREY)
                pdf.cell(MAIN_W - title_w, 4.5, dates, ln=1, align='R')

            elif len(parts) == 2:
                p0, p1 = parts[0], parts[1]
                import re as _re
                is_date = bool(_re.search(
                    r'\d{4}|Present|present|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', p1))
                if is_date:
                    pdf.set_x(MAIN_X + 2)
                    pdf.set_font("Arial", 'B', 9.5)
                    pdf.set_text_color(*DARK_GREY)
                    tw = min(pdf.get_string_width(p0) + 2, MAIN_W * 0.72)
                    pdf.cell(tw, 4.5, p0, ln=0)
                    pdf.set_font("Arial", 'I', 9)
                    pdf.set_text_color(*MID_GREY)
                    pdf.cell(MAIN_W - tw - 2, 4.5, p1, ln=1, align='R')
                else:
                    proj_name = p0
                    tech_stack = p1
                    pdf.set_font("Arial", 'B', 10)
                    name_w = pdf.get_string_width(proj_name) + 2
                    pdf.set_font("Arial", 'I', 9)
                    tech_w = pdf.get_string_width(tech_stack) + 2
                    if name_w + tech_w + 4 <= MAIN_W:
                        gap_w = MAIN_W - name_w - tech_w
                        pdf.set_font("Arial", 'B', 10)
                        pdf.set_text_color(*NAVY)
                        pdf.set_x(MAIN_X)
                        pdf.cell(name_w, 5, proj_name, ln=0)
                        pdf.cell(gap_w, 5, "", ln=0)
                        pdf.set_font("Arial", 'I', 9)
                        pdf.set_text_color(*MID_GREY)
                        pdf.cell(tech_w, 5, tech_stack, ln=1, align='R')
                    else:
                        pdf.set_font("Arial", 'B', 10)
                        pdf.set_text_color(*NAVY)
                        pdf.set_x(MAIN_X)
                        pdf.multi_cell(MAIN_W, 5, proj_name, align='L')
                        pdf.set_x(MAIN_X)
                        pdf.set_font("Arial", 'I', 9)
                        pdf.set_text_color(*MID_GREY)
                        pdf.multi_cell(MAIN_W, 4.5, tech_stack, align='L')
            else:
                pdf.set_font("Arial", 'B', 10)
                pdf.set_text_color(*NAVY)
                pdf.multi_cell(MAIN_W, 5, parts[0], align='L')

            pdf.set_text_color(*BLACK)
            continue

        if line.startswith("-"):
            pdf.set_x(MAIN_X + 3)
            pdf.set_font("Arial", size=9)
            pdf.set_text_color(*DARK_GREY)
            pdf.multi_cell(MAIN_W - 3, 4.5, "\x95 " + line[1:].lstrip(), align='L')
            continue

        pdf.set_x(MAIN_X)
        pdf.set_font("Arial", size=9)
        pdf.set_text_color(*DARK_GREY)
        pdf.multi_cell(MAIN_W, 4.5, line, align='L')


def _build_india_pdf(pdf: FPDF, text: str):
    """
    Jake Resume — strict 1-page ATS layout with auto font-size scaling.
    """
    MARGIN   = 10.0
    PAGE_W   = 210
    PAGE_H   = 297
    TEXT_W   = PAGE_W - 2 * MARGIN
    BOTTOM   = PAGE_H - MARGIN

    BLACK     = (0,   0,   0)
    DARK_GREY = (50,  50,  50)
    MID_GREY  = (110, 110, 110)

    SKIP_WORDS = {"NAME", "CONTACT", "SIDEBAR_START", "MAIN_START"}

    text  = text.replace("[SIDEBAR_START]", "").replace("[MAIN_START]", "")
    lines = [l.rstrip() for l in text.split("\n")]

    meaningful = sum(1 for l in lines if l.strip()
                     and l.strip() not in SKIP_WORDS
                     and l.strip() != "INTRODUCTION")
    if meaningful <= 52:
        BODY_PT   = 9.5;  HDR_PT = 10.5; LINE_H = 4.5; SEC_GAP = 5; BLK_GAP = 3
    elif meaningful <= 65:
        BODY_PT   = 9.0;  HDR_PT = 10.0; LINE_H = 4.2; SEC_GAP = 4; BLK_GAP = 2.5
    else:
        BODY_PT   = 8.5;  HDR_PT = 9.5;  LINE_H = 3.8; SEC_GAP = 3; BLK_GAP = 2

    BULLET_X = MARGIN + 3
    BULLET_W = TEXT_W - 3

    pdf.set_auto_page_break(auto=False)
    pdf.set_left_margin(MARGIN)
    pdf.set_right_margin(MARGIN)
    pdf.set_top_margin(MARGIN)
    pdf.set_y(MARGIN)

    name_line = ""
    for l in lines:
        s = l.strip()
        if s and s not in SKIP_WORDS and s != "INTRODUCTION" and "|" not in s and not s.startswith("-"):
            name_line = s
            break

    contact_line = ""
    for l in lines:
        s = l.strip()
        if "@" in s or s.startswith("+") or re.search(r"\d{7,}", s):
            if s != name_line:
                contact_line = s
                break

    name_printed    = False
    contact_printed = False
    intro_mode      = False
    in_experience   = False

    def fits(needed_mm=5):
        return pdf.get_y() + needed_mm < BOTTOM

    def draw_rule(thickness=0.3):
        if not fits(1): return
        pdf.set_draw_color(*BLACK)
        pdf.set_line_width(thickness)
        pdf.line(MARGIN, pdf.get_y(), MARGIN + TEXT_W, pdf.get_y())
        pdf.set_line_width(0.2)

    for raw_line in lines:
        line = raw_line.strip()

        if not fits(3):
            break

        if not line:
            if name_printed and fits(2):
                pdf.ln(0.8)
            continue

        if line in SKIP_WORDS:
            continue

        # NAME
        if not name_printed and line == name_line:
            display_name = line.title()
            pdf.set_font("Arial", "B", int(HDR_PT+7))
            pdf.set_text_color(*BLACK)
            pdf.set_x(MARGIN)
            pdf.cell(TEXT_W, 8, display_name, ln=True, align="C")
            name_printed = True
            continue

        # CONTACT
        if name_printed and not contact_printed and line == contact_line:
            pdf.set_font("Arial", "", BODY_PT-1)
            pdf.set_text_color(*DARK_GREY)
            pdf.set_x(MARGIN)
            pdf.multi_cell(TEXT_W, LINE_H-0.5, line, align="C")
            pdf.set_draw_color(*MID_GREY)
            pdf.set_line_width(0.2)
            pdf.line(MARGIN, pdf.get_y() + 0.5, MARGIN + TEXT_W, pdf.get_y() + 0.5)
            pdf.set_line_width(0.2)
            pdf.ln(BLK_GAP-0.5)
            contact_printed = True
            continue

        # INTRODUCTION keyword
        if line == "INTRODUCTION":
            intro_mode = True
            pdf.ln(1)
            continue

        # Intro paragraph
        if intro_mode:
            if (line.isupper() and len(line) > 3 and "|" not in line) or line.startswith("##"):
                intro_mode = False
            else:
                if fits(5):
                    pdf.set_x(MARGIN)
                    pdf.set_font("Arial", "I", BODY_PT-1)
                    pdf.set_text_color(*DARK_GREY)
                    pdf.multi_cell(TEXT_W, LINE_H-0.5, line, align="L")
                continue

        # ##COMPANY## HEADER
        if line.startswith("##COMPANY##"):
            company_name = line.replace("##COMPANY##", "").strip()
            if fits(30):
                pdf.ln(BLK_GAP)
            pdf.set_x(MARGIN)
            pdf.set_font("Arial", "B", HDR_PT)
            pdf.set_text_color(*BLACK)
            pdf.cell(TEXT_W, 5, company_name.upper(), ln=True)
            draw_rule(0.4)
            pdf.ln(0.8)
            continue

        # SECTION HEADER
        if (line.isupper()
                and 3 < len(line) <= 35
                and "|" not in line
                and line not in SKIP_WORDS
                and not any(c.isdigit() for c in line)):
            in_experience = ("EXPERIENCE" in line)
            if fits(30):
                pdf.ln(BLK_GAP)
            pdf.set_x(MARGIN)
            pdf.set_font("Arial", "B", HDR_PT-0.5)
            pdf.set_text_color(*BLACK)
            pdf.cell(TEXT_W, LINE_H, line, ln=True, align="L")
            draw_rule(0.3)
            pdf.ln(BLK_GAP-1)
            continue

        # ROLE / PROJECT PIPE LINE
        if "|" in line and not line.startswith("-"):
            parts = [p.strip() for p in line.split("|")]
            if fits(5):
                pdf.ln(BLK_GAP-1)

            if len(parts) >= 3:
                title, company, dates = parts[0], parts[1], parts[2]
                pdf.set_font("Arial", "B", HDR_PT-0.5)
                pdf.set_text_color(*BLACK)
                tw = min(pdf.get_string_width(title) + 2, TEXT_W * 0.74)
                pdf.set_x(MARGIN)
                pdf.cell(tw, 4.5, title, ln=0)
                pdf.set_font("Arial", "I", BODY_PT-0.5)
                pdf.set_text_color(*MID_GREY)
                pdf.cell(TEXT_W - tw, 4.5, dates, ln=1, align="R")
                pdf.set_x(MARGIN)
                pdf.set_font("Arial", "I", BODY_PT-0.5)
                pdf.set_text_color(*DARK_GREY)
                pdf.cell(TEXT_W, LINE_H-0.5, company, ln=True)

            elif len(parts) == 2:
                p0, p1 = parts[0], parts[1]
                is_date = bool(re.search(
                    r'\d{4}|Present|present|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', p1))
                if is_date:
                    pdf.set_font("Arial", "B", HDR_PT-0.5)
                    pdf.set_text_color(*BLACK)
                    tw = min(pdf.get_string_width(p0) + 2, TEXT_W * 0.74)
                    pdf.set_x(MARGIN + 2)
                    pdf.cell(tw, 4.5, p0, ln=0)
                    pdf.set_font("Arial", "I", BODY_PT-0.5)
                    pdf.set_text_color(*MID_GREY)
                    pdf.cell(TEXT_W - tw - 2, 4.5, p1, ln=1, align="R")
                else:
                    proj_name, tech_stack = p0, p1
                    pdf.set_font("Arial", "B", HDR_PT-0.5)
                    name_w = pdf.get_string_width(proj_name) + 2
                    pdf.set_font("Arial", "I", BODY_PT-0.5)
                    tech_w = pdf.get_string_width(tech_stack) + 2
                    pdf.set_x(MARGIN)
                    if name_w + tech_w + 4 <= TEXT_W:
                        pdf.set_font("Arial", "B", HDR_PT-0.5)
                        pdf.set_text_color(*BLACK)
                        pdf.cell(name_w, LINE_H, proj_name, ln=0)
                        pdf.cell(TEXT_W - name_w - tech_w, 4.5, "", ln=0)
                        pdf.set_font("Arial", "I", BODY_PT-0.5)
                        pdf.set_text_color(*MID_GREY)
                        pdf.cell(tech_w, LINE_H, tech_stack, ln=1, align="R")
                    else:
                        pdf.set_font("Arial", "B", HDR_PT-0.5)
                        pdf.set_text_color(*BLACK)
                        pdf.multi_cell(TEXT_W, LINE_H, proj_name, align="L")
                        pdf.set_x(MARGIN)
                        pdf.set_font("Arial", "I", BODY_PT-0.5)
                        pdf.set_text_color(*MID_GREY)
                        pdf.multi_cell(TEXT_W, LINE_H-0.5, tech_stack, align="L")
            else:
                pdf.set_font("Arial", "B", HDR_PT-0.5)
                pdf.set_text_color(*BLACK)
                pdf.set_x(MARGIN)
                pdf.cell(TEXT_W, 4.5, parts[0], ln=True)

            pdf.set_text_color(*BLACK)
            continue

        # SKILL LINE  "Category: items"
        if ":" in line and not line.startswith("-"):
            colon_idx = line.index(":")
            cat = line[:colon_idx].strip()
            det = line[colon_idx + 1:].strip()
            if cat and det and len(cat.split()) <= 5:
                pdf.set_x(MARGIN)
                pdf.set_font("Arial", "B", BODY_PT-0.5)
                pdf.set_text_color(*BLACK)
                lw = min(pdf.get_string_width(cat + ":  ") + 1, TEXT_W * 0.38)
                pdf.cell(lw, LINE_H-0.5, cat + ": ", ln=0)
                pdf.set_font("Arial", "", BODY_PT-0.5)
                pdf.multi_cell(TEXT_W - lw, LINE_H-0.5, det, align="L")
                continue

        # BULLET
        if line.startswith("- ") or line.startswith("-"):
            content = line[1:].lstrip()
            x_pos = BULLET_X + 2 if in_experience else BULLET_X
            w     = BULLET_W - 2 if in_experience else BULLET_W
            if fits(4):
                pdf.set_font("Arial", "", BODY_PT-0.5)
                pdf.set_text_color(*BLACK)
                pdf.set_x(x_pos)
                pdf.multi_cell(w, LINE_H-0.5, "\x95 " + content, align="L")
            continue

        # REGULAR TEXT
        if fits(4):
            pdf.set_font("Arial", "", BODY_PT-0.5)
            pdf.set_text_color(*BLACK)
            pdf.set_x(MARGIN)
            pdf.multi_cell(TEXT_W, LINE_H-0.5, line, align="L")


# ==============================================================================
# 5. UI
# ==============================================================================

st.markdown("""
<div class="main-header">
    <h1>🃏 BeTheJack</h1>
    <p>Upload your real CV → tailor it to any job, instantly.</p>
</div>
""", unsafe_allow_html=True)

ai_connected = init_ai()

for key, default in {
    "raw_cv_text": "",
    "tailored_content": "",
    "cv_filename": "",
    "photo_size": 52,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ──────────────────────────────────────────────
# STEP 1 + 2 — Upload & JD
# ──────────────────────────────────────────────
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("#### <span class='step-badge'>1</span> Upload Your CV", unsafe_allow_html=True)
    st.markdown('<div class="info-box">Supported: PDF, DOCX, TXT — your real CV, not a template.</div>', unsafe_allow_html=True)

    uploaded_cv = st.file_uploader(
        "Drop your CV here",
        type=["pdf", "docx", "doc", "txt"],
        label_visibility="collapsed"
    )

    if uploaded_cv:
        if uploaded_cv.name != st.session_state.cv_filename:
            with st.spinner("Reading your CV..."):
                extracted = parse_cv_file(uploaded_cv)
                st.session_state.raw_cv_text = extracted
                st.session_state.cv_filename = uploaded_cv.name
                st.session_state.tailored_content = ""

        if st.session_state.raw_cv_text:
            st.markdown('<div class="success-box">✅ CV parsed successfully!</div>', unsafe_allow_html=True)
            with st.expander("👁️ Preview extracted text"):
                st.text(st.session_state.raw_cv_text[:3000] + ("..." if len(st.session_state.raw_cv_text) > 3000 else ""))
        else:
            st.error("Could not extract text. Is the CV a scanned image? Try a text-based PDF or DOCX.")

    st.markdown("---")
    st.markdown("#### <span class='step-badge'>2</span> Select Layout", unsafe_allow_html=True)
    mode = st.radio(
        "Layout style",
        ["🌍 Global — Sidebar + Photo", "🇮🇳 ATS — Jake's Resume Style (Single Column)"],
        label_visibility="collapsed"
    )
    style_choice = "Global" if "Global" in mode else "India"

    uploaded_photo = None
    photo_size     = st.session_state.photo_size
    if style_choice == "Global":
        uploaded_photo = st.file_uploader("Upload Profile Photo (optional)", type=['jpg', 'jpeg', 'png'])
        if uploaded_photo:
            photo_size = st.slider(
                "📐 Photo size (mm)",
                min_value=35, max_value=68, value=st.session_state.photo_size, step=1,
                help="Adjust how large your profile photo appears in the sidebar"
            )
            st.session_state.photo_size = photo_size


with col2:
    st.markdown("#### <span class='step-badge'>3</span> Paste Job Description", unsafe_allow_html=True)
    job_desc = st.text_area(
        "Job Description",
        height=340,
        placeholder="Paste the full job description here...\n\nThe more detail you provide, the better the tailoring.",
        label_visibility="collapsed"
    )

# ──────────────────────────────────────────────
# GENERATE BUTTON
# ──────────────────────────────────────────────
st.markdown("---")
col_btn, col_info = st.columns([1, 2])
with col_btn:
    generate_btn = st.button("✨ Tailor My CV", type="primary", use_container_width=True)

with col_info:
    st.markdown("""
    <div class="info-box">
    <b>What this does:</b> Rewrites your real experience using JD keywords, reorders skills by relevance,
    and crafts a targeted summary — <em>all based only on what's already in your CV</em>.
    </div>
    """, unsafe_allow_html=True)

if generate_btn:
    if not st.session_state.raw_cv_text:
        st.error("Please upload your CV first.")
    elif not job_desc.strip():
        st.error("Please paste a Job Description.")
    elif not ai_connected:
        st.error("AI not connected. Check your API key in Secrets.")
    else:
        with st.spinner("Tailoring your CV to the job description... this takes ~20 seconds."):
            result = tailor_cv(
                st.session_state.raw_cv_text,
                job_desc,
                style=style_choice
            )
            if result.startswith("Error generating content:"):
                err = result
                if "429" in err or "rate" in err.lower() or "quota" in err.lower():
                    st.error("⚠️ Groq API rate limit hit. Wait a moment and try again, or check your quota at console.groq.com")
                else:
                    st.error(err)
            else:
                st.session_state.tailored_content = result

# ──────────────────────────────────────────────
# EDIT + RENDER
# ──────────────────────────────────────────────
if st.session_state.tailored_content:
    st.markdown("---")
    st.markdown("#### <span class='step-badge'>4</span> Review & Edit the Tailored Draft", unsafe_allow_html=True)
    st.markdown('<div class="warning-box">⚠️ Always review before downloading. Check dates, titles, and facts are still accurate.</div>', unsafe_allow_html=True)

    edited = st.text_area(
        "Tailored CV Content",
        value=st.session_state.tailored_content,
        height=580,
        label_visibility="collapsed"
    )
    st.session_state.tailored_content = edited

    col_pdf, col_docx, col_spacer = st.columns([1, 1, 1])
    with col_pdf:
        render_btn = st.button("📄 Render PDF", type="secondary", use_container_width=True)
    with col_docx:
        docx_btn = st.button("📝 Download DOCX", type="secondary", use_container_width=True)

    if render_btn:
        with st.spinner("Building your PDF..."):
            photo_path = None
            if uploaded_photo:
                photo_path = "temp_profile_photo.jpg"
                with open(photo_path, "wb") as f:
                    f.write(uploaded_photo.getbuffer())

            pdf_bytes = build_pdf(
                st.session_state.tailored_content,
                style_choice,
                photo_path=photo_path,
                photo_size=st.session_state.photo_size
            )

            if photo_path and os.path.exists(photo_path):
                try: os.remove(photo_path)
                except: pass

            safe_title = re.sub(r'[^a-zA-Z0-9]', '_', job_desc[:25]) if job_desc else "Resume"
            filename = f"CV_{style_choice}_{safe_title}.pdf"

        st.success("✅ PDF is ready!")
        st.download_button(
            label="📥 Download PDF",
            data=pdf_bytes,
            file_name=filename,
            mime="application/pdf",
            use_container_width=True
        )

    if docx_btn:
        with st.spinner("Building your DOCX..."):
            try:
                docx_bytes = build_docx(st.session_state.tailored_content)
                safe_title = re.sub(r'[^a-zA-Z0-9]', '_', job_desc[:25]) if job_desc else "Resume"
                docx_filename = f"CV_{safe_title}.docx"
                st.success("✅ DOCX is ready!")
                st.download_button(
                    label="📥 Download DOCX",
                    data=docx_bytes,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align:center;color:#888;font-size:0.8rem;'>BeTheJack · Smart CV tailoring, powered by AI.</div>",
    unsafe_allow_html=True
)
